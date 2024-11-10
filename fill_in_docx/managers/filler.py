import logging
from docx import Document

# Налаштування логування
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)


class TemplateFiller:
    """Заповнює шаблон договору та редагує зазначені абзаци"""

    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path
        logging.info(
            "TemplateFiller створений з шляхом до шаблону: %s і вихідним шляхом: %s",
            template_path,
            output_path,
        )

    def fill_and_edit(
        self,
        data: dict,
        text_to_delete: list[str] = None,
        bold_keys: list[str] = None,
    ):
        """Заповнює шаблон договору та видаляє зазначені абзаци"""

        if bold_keys is None:
            bold_keys = []
        logging.info(
            "Початок обробки шаблону. Кількість маркерів для виділення жирним: %d",
            len(bold_keys),
        )

        try:
            # Завантаження шаблону
            doc = Document(self.template_path)
            logging.info("Шаблон завантажено успішно з %s", self.template_path)

            # Видалення абзаців з вказаним текстом
            if text_to_delete:
                self._delete_paragraph(doc, text_to_delete)

            # Заповнення шаблону
            self._process_paragraphs(doc, data, bold_keys)
            self._process_tables(doc, data, bold_keys)

            # Збереження заповненого документа
            doc.save(self.output_path)
            logging.info("Документ збережено успішно в %s", self.output_path)

        except Exception as e:
            logging.error("Помилка під час обробки шаблону: %s", str(e))

    @staticmethod
    def _delete_paragraph(doc, texts_to_delete: list = None):
        """Видаляє абзаци, що містять будь-який з вказаних текстів, з документа"""

        if texts_to_delete is None:
            texts_to_delete = []
        logging.info(
            "Початок видалення абзаців з текстами, кількість: %d",
            len(texts_to_delete),
        )

        logging.info("Видалення абзаців з текстами: %s", texts_to_delete)
        for paragraph in doc.paragraphs:
            for text in texts_to_delete:
                if text in paragraph.text:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                    p_element.clear()
                    logging.info("Абзац з текстом '%s' видалено.", text)
                    break  # Вийти з циклу text після видалення абзацу

    @staticmethod
    def _replace_markers_in_text(
        text: str, marker_key: str, marker_value: str, is_bold: bool
    ):
        """Замінює маркери в тексті на відповідні значення"""

        if f"{{{marker_key}}}" in text:
            text = text.replace(f"{{{marker_key}}}", marker_value)
            return text, is_bold
        return text, False

    def _process_paragraphs(self, doc, data, bold_keys):
        """Обробляє абзаци документа для заміни маркерів"""

        logging.info("Обробка абзаців документа...")
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run_text = run.text
                for key, value in data.items():
                    run_text, should_bold = self._replace_markers_in_text(
                        run_text, key, value, key in bold_keys
                    )
                    run.text = run_text
                    if should_bold:
                        run.bold = True
        logging.info("Обробка абзаців завершена.")

    def _process_tables(self, doc, data, bold_keys):
        """Обробляє таблиці документа для заміни маркерів"""

        logging.info("Обробка таблиць документа...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            for key, value in data.items():
                                run_text, should_bold = self._replace_markers_in_text(
                                    run_text, key, value, key in bold_keys
                                )
                                run.text = run_text
                                if should_bold:
                                    run.bold = True
        logging.info("Обробка таблиць завершена.")
